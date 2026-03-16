"""Generates a tailored DOCX resume from master resume text + AI suggestions."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from typing import Optional

RESUMES_DIR = Path(__file__).parent.parent / "resumes"
OUTPUTS_DIR = Path(__file__).parent.parent / "outputs"


def load_master_resume() -> str:
    """Read master resume text from resumes/ directory."""
    RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    for ext in ("*.txt", "*.md"):
        files = list(RESUMES_DIR.glob(ext))
        if files:
            return files[0].read_text()
    return ""


def save_tailored_cv(company: str, role: str, content: str) -> Path:
    """Write tailored resume as a DOCX file and return its path."""
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^\w\-]", "_", f"{company}_{role}")[:60]
    out_path = OUTPUTS_DIR / f"CV_{safe_name}.docx"

    doc = Document()
    doc.add_heading(role, 0)
    doc.add_paragraph(f"Tailored for: {company}")
    doc.add_paragraph("")

    for line in content.split("\n"):
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(out_path)
    return out_path
